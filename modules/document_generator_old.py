"""
PDF and Word Document Generator Module
Generates professional quotation documents in PDF and Word formats
"""

import os
from datetime import datetime
from xhtml2pdf import pisa
from io import BytesIO
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from typing import Dict, Any


class DocumentGenerator:
    """Generate quotation documents in PDF and Word formats"""
    
    def __init__(self, output_folder: str = 'outputs'):
        self.output_folder = output_folder
        os.makedirs(output_folder, exist_ok=True)
    
    def generate_pdf(self, quotation_data: Dict[str, Any], template_type: str) -> str:
        """
        Generate PDF from quotation data
        
        Args:
            quotation_data: Dictionary containing quotation information
            template_type: 'type1' or 'type2'
            
        Returns:
            Path to generated PDF file
        """
        html_content = self._create_html_content(quotation_data, template_type)
        
        # Generate filename
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        filename = f"quotation_{template_type}_{timestamp}.pdf"
        output_path = os.path.join(self.output_folder, filename)
        
        # Convert HTML to PDF using xhtml2pdf
        with open(output_path, 'wb') as pdf_file:
            pisa_status = pisa.CreatePDF(
                html_content,
                dest=pdf_file
            )
        
        if pisa_status.err:
            raise Exception(f"PDF generation failed with error code: {pisa_status.err}")
        
        return output_path
    
    def generate_word(self, quotation_data: Dict[str, Any], template_type: str) -> str:
        """
        Generate Word document from quotation data
        
        Args:
            quotation_data: Dictionary containing quotation information
            template_type: 'type1' or 'type2'
            
        Returns:
            Path to generated Word file
        """
        doc = Document()
        
        # Set document margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.75)
            section.bottom_margin = Inches(0.75)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        
        if template_type == 'type1':
            self._create_word_type1(doc, quotation_data)
        else:
            self._create_word_type2(doc, quotation_data)
        
        # Generate filename
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        filename = f"quotation_{template_type}_{timestamp}.docx"
        output_path = os.path.join(self.output_folder, filename)
        
        doc.save(output_path)
        
        return output_path
    
    def _create_html_content(self, data: Dict[str, Any], template_type: str) -> str:
        """Create HTML content for PDF generation"""
        
        if template_type == 'type1':
            return self._create_html_type1(data)
        else:
            return self._create_html_type2(data)
    
    def _create_html_type1(self, data: Dict[str, Any]) -> str:
        """Create HTML for Type 1 template (Detailed Itemized)"""
        
        # Build pricing table HTML
        pricing_rows = ""
        for item in data.get('pricing_table', []):
            pricing_rows += f"""
            <tr>
                <td>{item.get('item_no', '')}</td>
                <td>{item.get('description', '')}</td>
                <td>{item.get('quantity', '')}</td>
                <td>₹{item.get('unit_price', '')}</td>
                <td>₹{item.get('total_price', '')}</td>
            </tr>
            """
        
        # Build timeline table
        timeline_rows = ""
        for phase in data.get('timeline', []):
            timeline_rows += f"""
            <tr>
                <td>{phase.get('phase', '')}</td>
                <td>{phase.get('duration', '')}</td>
                <td>{phase.get('deliverables', '')}</td>
            </tr>
            """
        
        # Build scope of work
        scope_items = ""
        for item in data.get('scope_of_work', []):
            scope_items += f"<li>{item}</li>"
        
        # Build terms and conditions
        terms_items = ""
        for term in data.get('terms_and_conditions', []):
            terms_items += f"<li>{term}</li>"
        
        html = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="UTF-8">
            <style>
                @page {{
                    size: A4;
                    margin: 2cm;
                }}
                body {{
                    font-family: 'Segoe UI', Arial, sans-serif;
                    line-height: 1.6;
                    color: #333;
                }}
                .header {{
                    text-align: center;
                    border-bottom: 3px solid #2c3e50;
                    padding-bottom: 20px;
                    margin-bottom: 30px;
                }}
                .header h1 {{
                    color: #2c3e50;
                    margin: 0;
                    font-size: 28px;
                }}
                .info-section {{
                    display: flex;
                    justify-content: space-between;
                    margin-bottom: 30px;
                }}
                .info-box {{
                    flex: 1;
                }}
                h2 {{
                    color: #2c3e50;
                    border-bottom: 2px solid #3498db;
                    padding-bottom: 5px;
                    margin-top: 25px;
                }}
                table {{
                    width: 100%;
                    border-collapse: collapse;
                    margin: 15px 0;
                }}
                th {{
                    background-color: #2c3e50;
                    color: white;
                    padding: 12px;
                    text-align: left;
                    font-weight: bold;
                }}
                td {{
                    padding: 10px;
                    border-bottom: 1px solid #ddd;
                }}
                tr:nth-child(even) {{
                    background-color: #f8f9fa;
                }}
                .total-section {{
                    margin-top: 20px;
                    text-align: right;
                }}
                .total-row {{
                    display: flex;
                    justify-content: flex-end;
                    margin: 5px 0;
                }}
                .total-label {{
                    width: 150px;
                    font-weight: bold;
                }}
                .total-value {{
                    width: 150px;
                    text-align: right;
                }}
                .grand-total {{
                    font-size: 18px;
                    color: #2c3e50;
                    border-top: 2px solid #2c3e50;
                    padding-top: 10px;
                    margin-top: 10px;
                }}
                ul {{
                    line-height: 1.8;
                }}
                .footer {{
                    margin-top: 40px;
                    padding-top: 20px;
                    border-top: 2px solid #2c3e50;
                    text-align: center;
                    font-size: 12px;
                    color: #666;
                }}
            </style>
        </head>
        <body>
            <div class="header">
                <h1>QUOTATION PROPOSAL</h1>
                <p style="margin: 5px 0;">Detailed Itemized Quotation</p>
            </div>
            
            <div class="info-section">
                <div class="info-box">
                    <p><strong>Client:</strong> {data.get('client_name', 'Valued Client')}</p>
                    <p><strong>Project:</strong> {data.get('project_title', 'Project Quotation')}</p>
                </div>
                <div class="info-box" style="text-align: right;">
                    <p><strong>Date:</strong> {data.get('date', datetime.now().strftime('%d %B %Y'))}</p>
                    <p><strong>Reference:</strong> {data.get('reference_number', 'QT-2024-001')}</p>
                </div>
            </div>
            
            <h2>Executive Summary</h2>
            <p>{data.get('executive_summary', '')}</p>
            
            <h2>Scope of Work</h2>
            <ul>
                {scope_items}
            </ul>
            
            <h2>Detailed Pricing</h2>
            <table>
                <thead>
                    <tr>
                        <th>Item No.</th>
                        <th>Description</th>
                        <th>Quantity</th>
                        <th>Unit Price</th>
                        <th>Total Price</th>
                    </tr>
                </thead>
                <tbody>
                    {pricing_rows}
                </tbody>
            </table>
            
            <div class="total-section">
                <div class="total-row">
                    <div class="total-label">Subtotal:</div>
                    <div class="total-value">₹{data.get('subtotal', '0')}</div>
                </div>
                <div class="total-row">
                    <div class="total-label">Tax (GST):</div>
                    <div class="total-value">₹{data.get('tax', '0')}</div>
                </div>
                <div class="total-row grand-total">
                    <div class="total-label">Grand Total:</div>
                    <div class="total-value">₹{data.get('grand_total', '0')}</div>
                </div>
            </div>
            
            <h2>Implementation Timeline</h2>
            <table>
                <thead>
                    <tr>
                        <th>Phase</th>
                        <th>Duration</th>
                        <th>Deliverables</th>
                    </tr>
                </thead>
                <tbody>
                    {timeline_rows}
                </tbody>
            </table>
            
            <h2>Terms and Conditions</h2>
            <ul>
                {terms_items}
            </ul>
            
            <div class="footer">
                <p>This quotation is valid for 30 days from the date of issue.</p>
                <p>Thank you for considering our proposal.</p>
            </div>
        </body>
        </html>
        """
        
        return html
    
    def _create_html_type2(self, data: Dict[str, Any]) -> str:
        """Create HTML for Type 2 template (Executive Summary)"""
        
        # Build investment table
        pricing_rows = ""
        for item in data.get('pricing_table', []):
            pricing_rows += f"""
            <tr>
                <td>{item.get('description', '')}</td>
                <td>₹{item.get('total_price', '')}</td>
            </tr>
            """
        
        # Build scope items
        scope_items = ""
        for item in data.get('scope_of_work', []):
            scope_items += f"<li>{item}</li>"
        
        html = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="UTF-8">
            <style>
                @page {{
                    size: A4;
                    margin: 2cm;
                }}
                body {{
                    font-family: 'Segoe UI', Arial, sans-serif;
                    line-height: 1.6;
                    color: #333;
                }}
                .cover {{
                    text-align: center;
                    padding: 60px 0;
                    background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
                    color: white;
                    margin: -2cm -2cm 2cm -2cm;
                    padding: 80px 2cm;
                }}
                .cover h1 {{
                    font-size: 36px;
                    margin: 0 0 20px 0;
                }}
                .cover p {{
                    font-size: 18px;
                    margin: 10px 0;
                }}
                h2 {{
                    color: #667eea;
                    border-bottom: 3px solid #667eea;
                    padding-bottom: 8px;
                    margin-top: 30px;
                }}
                .executive-summary {{
                    background-color: #f8f9fa;
                    padding: 20px;
                    border-left: 4px solid #667eea;
                    margin: 20px 0;
                }}
                table {{
                    width: 100%;
                    border-collapse: collapse;
                    margin: 20px 0;
                }}
                th {{
                    background-color: #667eea;
                    color: white;
                    padding: 15px;
                    text-align: left;
                    font-weight: bold;
                }}
                td {{
                    padding: 12px 15px;
                    border-bottom: 1px solid #ddd;
                }}
                tr:nth-child(even) {{
                    background-color: #f8f9fa;
                }}
                .investment-total {{
                    background-color: #667eea !important;
                    color: white !important;
                    font-weight: bold;
                    font-size: 18px;
                }}
                ul {{
                    line-height: 2;
                }}
                .highlight-box {{
                    background-color: #f0f4ff;
                    padding: 20px;
                    border-radius: 8px;
                    margin: 20px 0;
                }}
                .footer {{
                    margin-top: 50px;
                    padding-top: 20px;
                    border-top: 3px solid #667eea;
                    text-align: center;
                }}
            </style>
        </head>
        <body>
            <div class="cover">
                <h1>{data.get('project_title', 'Project Proposal')}</h1>
                <p>Prepared for: {data.get('client_name', 'Valued Client')}</p>
                <p>{data.get('date', datetime.now().strftime('%d %B %Y'))}</p>
                <p>Reference: {data.get('reference_number', 'QT-2024-001')}</p>
            </div>
            
            <h2>Executive Overview</h2>
            <div class="executive-summary">
                <p>{data.get('executive_summary', '')}</p>
            </div>
            
            <h2>Solution Package</h2>
            <div class="highlight-box">
                <h3 style="margin-top: 0; color: #667eea;">Key Deliverables</h3>
                <ul>
                    {scope_items}
                </ul>
            </div>
            
            <h2>Investment Summary</h2>
            <table>
                <thead>
                    <tr>
                        <th>Component</th>
                        <th>Investment</th>
                    </tr>
                </thead>
                <tbody>
                    {pricing_rows}
                    <tr class="investment-total">
                        <td>Total Investment</td>
                        <td>₹{data.get('grand_total', '0')}</td>
                    </tr>
                </tbody>
            </table>
            
            <h2>Implementation Approach</h2>
            <p>{data.get('additional_notes', 'Our proven methodology ensures successful project delivery with regular milestones and quality checkpoints.')}</p>
            
            <h2>Next Steps</h2>
            <div class="highlight-box">
                <ol>
                    <li>Review and approve this proposal</li>
                    <li>Sign the agreement</li>
                    <li>Project kickoff within 5 business days</li>
                    <li>Regular progress updates throughout implementation</li>
                </ol>
            </div>
            
            <div class="footer">
                <p><strong>This proposal is valid for 30 days from the date of issue.</strong></p>
                <p>We look forward to partnering with you on this exciting project.</p>
            </div>
        </body>
        </html>
        """
        
        return html
    
    def _create_word_type1(self, doc: Document, data: Dict[str, Any]):
        """Create Word document for Type 1 template"""
        
        # Title
        title = doc.add_heading('QUOTATION PROPOSAL', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        subtitle = doc.add_paragraph('Detailed Itemized Quotation')
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
        
        # Client info
        info = doc.add_paragraph()
        info.add_run(f"Client: ").bold = True
        info.add_run(data.get('client_name', 'Valued Client'))
        info.add_run('\n')
        info.add_run(f"Project: ").bold = True
        info.add_run(data.get('project_title', 'Project Quotation'))
        info.add_run('\n')
        info.add_run(f"Date: ").bold = True
        info.add_run(data.get('date', datetime.now().strftime('%d %B %Y')))
        info.add_run('\n')
        info.add_run(f"Reference: ").bold = True
        info.add_run(data.get('reference_number', 'QT-2024-001'))
        
        # Executive Summary
        doc.add_heading('Executive Summary', 1)
        doc.add_paragraph(data.get('executive_summary', ''))
        
        # Scope of Work
        doc.add_heading('Scope of Work', 1)
        for item in data.get('scope_of_work', []):
            doc.add_paragraph(item, style='List Bullet')
        
        # Pricing Table
        doc.add_heading('Detailed Pricing', 1)
        pricing_table = data.get('pricing_table', [])
        if pricing_table:
            table = doc.add_table(rows=1, cols=5)
            table.style = 'Light Grid Accent 1'
            
            # Header row
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Item No.'
            hdr_cells[1].text = 'Description'
            hdr_cells[2].text = 'Quantity'
            hdr_cells[3].text = 'Unit Price'
            hdr_cells[4].text = 'Total Price'
            
            # Data rows
            for item in pricing_table:
                row_cells = table.add_row().cells
                row_cells[0].text = str(item.get('item_no', ''))
                row_cells[1].text = str(item.get('description', ''))
                row_cells[2].text = str(item.get('quantity', ''))
                row_cells[3].text = f"₹{item.get('unit_price', '')}"
                row_cells[4].text = f"₹{item.get('total_price', '')}"
        
        # Totals
        doc.add_paragraph()
        totals = doc.add_paragraph()
        totals.add_run(f"Subtotal: ₹{data.get('subtotal', '0')}").bold = True
        totals.add_run('\n')
        totals.add_run(f"Tax (GST): ₹{data.get('tax', '0')}").bold = True
        totals.add_run('\n')
        totals.add_run(f"Grand Total: ₹{data.get('grand_total', '0')}").bold = True
        totals.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Timeline
        doc.add_heading('Implementation Timeline', 1)
        timeline = data.get('timeline', [])
        if timeline:
            table = doc.add_table(rows=1, cols=3)
            table.style = 'Light Grid Accent 1'
            
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Phase'
            hdr_cells[1].text = 'Duration'
            hdr_cells[2].text = 'Deliverables'
            
            for phase in timeline:
                row_cells = table.add_row().cells
                row_cells[0].text = str(phase.get('phase', ''))
                row_cells[1].text = str(phase.get('duration', ''))
                row_cells[2].text = str(phase.get('deliverables', ''))
        
        # Terms
        doc.add_heading('Terms and Conditions', 1)
        for term in data.get('terms_and_conditions', []):
            doc.add_paragraph(term, style='List Bullet')
    
    def _create_word_type2(self, doc: Document, data: Dict[str, Any]):
        """Create Word document for Type 2 template"""
        
        # Cover page
        title = doc.add_heading(data.get('project_title', 'Project Proposal'), 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
        
        cover_info = doc.add_paragraph()
        cover_info.add_run(f"Prepared for: {data.get('client_name', 'Valued Client')}").bold = True
        cover_info.add_run(f"\n{data.get('date', datetime.now().strftime('%d %B %Y'))}")
        cover_info.add_run(f"\nReference: {data.get('reference_number', 'QT-2024-001')}")
        cover_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_page_break()
        
        # Executive Overview
        doc.add_heading('Executive Overview', 1)
        doc.add_paragraph(data.get('executive_summary', ''))
        
        # Solution Package
        doc.add_heading('Solution Package - Key Deliverables', 1)
        for item in data.get('scope_of_work', []):
            doc.add_paragraph(item, style='List Bullet')
        
        # Investment Summary
        doc.add_heading('Investment Summary', 1)
        pricing_table = data.get('pricing_table', [])
        if pricing_table:
            table = doc.add_table(rows=1, cols=2)
            table.style = 'Light Grid Accent 1'
            
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Component'
            hdr_cells[1].text = 'Investment'
            
            for item in pricing_table:
                row_cells = table.add_row().cells
                row_cells[0].text = str(item.get('description', ''))
                row_cells[1].text = f"₹{item.get('total_price', '')}"
            
            # Total row
            total_cells = table.add_row().cells
            total_cells[0].text = 'Total Investment'
            total_cells[1].text = f"₹{data.get('grand_total', '0')}"
            for cell in total_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
        
        # Next Steps
        doc.add_heading('Next Steps', 1)
        steps = [
            'Review and approve this proposal',
            'Sign the agreement',
            'Project kickoff within 5 business days',
            'Regular progress updates throughout implementation'
        ]
        for step in steps:
            doc.add_paragraph(step, style='List Number')
